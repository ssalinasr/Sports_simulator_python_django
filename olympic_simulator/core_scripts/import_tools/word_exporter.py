"""Generación del informe olímpico en formato Word.

El módulo recibe la estructura ``sports_dict`` construida en
``appolympics.views.exportar_word``. No consulta la base de datos: los títulos
de las pruebas se obtienen de ``ligas_deporte`` y sus podios de
``tablas_resultados``, respetando el mismo orden.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import Any, Iterable, Mapping, Optional, Union

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


FONT_NAME = "Arial Narrow"


def _color_rgb(color: Any) -> RGBColor:
    """Convierte colores hexadecimales (``#RRGGBB``) a ``RGBColor``.

    Un color inválido no debe impedir la exportación; en ese caso se usa negro.
    """

    value = str(color or "").strip().lstrip("#")
    rgb_match = re.fullmatch(r"rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)", value, re.I)
    if rgb_match:
        channels = [min(255, int(channel)) for channel in rgb_match.groups()]
        return RGBColor(*channels)
    if re.fullmatch(r"[0-9A-Fa-f]{3}", value):
        value = "".join(channel * 2 for channel in value)
    if re.fullmatch(r"[0-9A-Fa-f]{6}", value):
        return RGBColor.from_string(value.upper())
    return RGBColor(0, 0, 0)


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, uppercase: bool = False) -> None:
    """Escribe una celda aplicando la fuente requerida a todos sus runs."""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("" if text is None else str(text).upper() if uppercase else str(text))
    run.bold = bold
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_borders(table: Any) -> None:
    """Añade bordes visibles a una tabla de Word."""

    table_properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    table_properties.append(borders)


def _add_title(document: Document, text: Any, color: Any, *, size: int, centered: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(str(text))
    run.bold = True
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = _color_rgb(color)


def _attribute_or_value(value: Any, attribute: str, default: str = "") -> str:
    """Obtiene un atributo de un modelo u homónimo de un diccionario."""

    if isinstance(value, Mapping):
        return str(value.get(attribute, default))
    return str(getattr(value, attribute, default))


def _sport_name(sport_data: Mapping[str, Any]) -> str:
    name = sport_data.get("nombre_deporte")
    if name:
        return str(name)
    tournaments = sport_data.get("torneos_deporte", [])
    if tournaments:
        return _attribute_or_value(tournaments[0], "team_sport_name", "Deporte")
    return "Deporte"


def _discipline_name(discipline: Any, position: int) -> str:
    return _attribute_or_value(discipline, "sp_record_name", f"Prueba {position}")


def _podium_values(row: Any) -> tuple[Any, Any, Any]:
    """Extrae representante, país y resultado de una fila del diccionario."""

    if isinstance(row, Mapping):
        return (
            row.get("representante", row.get("participant", "---")),
            row.get("pais", row.get("country", "---")),
            row.get("resultado", row.get("result", "---")),
        )
    values = list(row) if isinstance(row, (tuple, list)) else []
    # Las filas generadas por la vista tienen: medalla, representante, país, resultado.
    if len(values) >= 4:
        return values[1], values[2], values[3]
    if len(values) == 3:
        return values[1], values[2], "---"
    return "---", "---", "---"


def _medal_rows(medallero: Any) -> Iterable[tuple[Any, Mapping[str, Any]]]:
    """Admite el listado ordenado de la vista o un diccionario país -> medallas."""

    if isinstance(medallero, Mapping):
        return medallero.items()
    return medallero or []


def _add_podium_table(document: Document, results: Any) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    _set_table_borders(table)

    for cell, header in zip(table.rows[0].cells, ("REPRESENTANTE", "PAÍS", "RESULTADO")):
        _set_cell_text(cell, header, bold=True, uppercase=True)

    result_rows = results if isinstance(results, (list, tuple)) else []
    for position in range(1, 4):
        representative, country, result = _podium_values(
            result_rows[position - 1] if position <= len(result_rows) else None
        )
        cells = table.add_row().cells
        _set_cell_text(cells[0], f"{position}. {representative}")
        _set_cell_text(cells[1], country)
        _set_cell_text(cells[2], result)


def _add_medal_table(document: Document, medallero: Any) -> None:
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    _set_table_borders(table)

    for cell, header in zip(table.rows[0].cells, ("PAÍS", "ORO", "PLATA", "BRONCE", "TOTAL")):
        _set_cell_text(cell, header, bold=True, uppercase=True)

    for country, medals in _medal_rows(medallero):
        medals = medals if isinstance(medals, Mapping) else {}
        cells = table.add_row().cells
        _set_cell_text(cells[0], country)
        _set_cell_text(cells[1], medals.get("O", medals.get("oro", 0)))
        _set_cell_text(cells[2], medals.get("P", medals.get("plata", 0)))
        _set_cell_text(cells[3], medals.get("B", medals.get("bronce", 0)))
        _set_cell_text(cells[4], medals.get("Total", medals.get("total", 0)))


def exportar_word(
    sports_dict: Mapping[str, Mapping[str, Any]],
    year: Any,
    medallero_general: Any,
    output_path: Optional[Union[str, Path]] = None,
) -> Union[BytesIO, Path]:
    """Crea el reporte Word de los resultados olímpicos.

    Args:
        sports_dict: Diccionario construido en la vista, con las claves
            ``ligas_deporte``, ``tablas_resultados``, ``medallero_deporte`` y
            ``color_deporte`` por deporte.
        year: Año que se muestra en el título principal.
        medallero_general: Listado ``[(pais, {O, P, B, Total}), ...]`` o
            diccionario equivalente para la tabla final.
        output_path: Ruta opcional del ``.docx``. Si no se indica, devuelve un
            ``BytesIO`` listo para usarse con ``FileResponse``.
    """

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal_style.font.size = Pt(10)

    _add_title(document, str(year), "000000", size=24, centered=True)

    for sport_data in sports_dict.values():
        color = sport_data.get("color_deporte", "000000")
        _add_title(document, _sport_name(sport_data), color, size=16)

        disciplines = sport_data.get("ligas_deporte", [])
        result_tables = sport_data.get("tablas_resultados", [])
        for index, results in enumerate(result_tables, start=1):
            discipline = disciplines[index - 1] if index <= len(disciplines) else None
            _add_title(document, _discipline_name(discipline, index), color, size=12)
            _add_podium_table(document, results)

        _add_title(document, f"Medallero de {_sport_name(sport_data)}", color, size=12)
        _add_medal_table(document, sport_data.get("medallero_deporte", []))

    _add_title(document, "Resultados Finales", "000000", size=16)
    _add_medal_table(document, medallero_general)

    if output_path is not None:
        destination = Path(output_path)
        document.save(str(destination))
        return destination

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


# Alias semántico para usos que prefieren el verbo "generar" en la vista.
generar_word = exportar_word
